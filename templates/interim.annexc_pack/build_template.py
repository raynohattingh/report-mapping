"""Build pack_template.docx (docxtpl template) — run once, artifact is committed.

The template is DATA (Constitution IV): revising the pack layout means editing
this script or the .docx directly, never pipeline code. INTERIM stand-in only
(Constitution I, A2): structure modeled on DST 34-1441 Annex C, zero invented
Eskom content.

Usage: uv run python templates/interim.annexc_pack/build_template.py
"""

from pathlib import Path

from docx import Document


def build() -> Path:
    doc = Document()
    doc.add_heading("Structure Inspection Report Pack (INTERIM)", level=0)
    doc.add_paragraph(
        "INTERIM stand-in template. The client-mandated pro forma (TBD-1) replaces "
        "this as a new TargetTemplate version when obtained."
    )

    doc.add_heading("Inspection details", level=1)
    details = doc.add_table(rows=5, cols=2)
    labels = [
        ("Inspection name", "{{ inspection_name }}"),
        ("Inspection date", "{{ inspection_date }}"),
        ("Contract number", "{{ contract_number }}"),
        ("Inspection method", "{{ inspection_method }}"),
        ("Company", "{{ company }}"),
    ]
    for row, (label, value) in zip(details.rows, labels):
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_heading("Findings", level=1)
    table = doc.add_table(rows=3, cols=5)
    header = table.rows[0].cells
    for i, name in enumerate(["Finding ID", "Defect code", "Priority", "Severity", "Comments"]):
        header[i].text = name
    table.rows[1].cells[0].text = "{%tr for f in findings %}"
    body = table.rows[2].cells
    for i, field in enumerate(
        ["{{ f.finding_id }}", "{{ f.defect_code }}", "{{ f.priority }}",
         "{{ f.source_severity }}", "{{ f.comments }}"]
    ):
        body[i].text = field
    # docxtpl consumes the {%tr %} rows; append the endfor row.
    endrow = table.add_row()
    endrow.cells[0].text = "{%tr endfor %}"

    doc.add_paragraph("Exceptions for this report are listed in the batch exceptions report.")

    out = Path(__file__).parent / "pack_template.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build())
